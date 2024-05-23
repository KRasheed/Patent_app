from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

document = Document()


# Set default font and size
def default_format():
    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)  # Black color
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


# Function to apply required style to headings
def add_heading_with_style(text, level=3):
    heading = document.add_heading(level=level)
    run = heading.add_run(text)
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)
    paragraph_format = heading.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return heading


def add_subheading_with_style(text, level=4):
    subheading = document.add_heading(level=level)
    run = subheading.add_run(text)
    run.font.italic = True
    run.font.bold = False
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)
    paragraph_format = subheading.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return subheading


# Function to add paragraph with justified alignment
def add_paragraph_with_style(text, numbered=False):
    if numbered:
        paragraph = document.add_paragraph(style="List Number")
    else:
        paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)  # Ensure text is black
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return paragraph


# Function to set margins
def set_margins(doc, top, right, bottom, left):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.right_margin = Inches(right)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
